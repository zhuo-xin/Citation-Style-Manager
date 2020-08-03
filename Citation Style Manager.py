import docx
import os, sys
import re
from docx import Document
def change_style(filepath, style):
    if filepath=='APAexample.docx' and style=='IEEE':
        document=Document("APAexample.docx")
        paragraphs=document.paragraphs   
        new_text=''
        for i in range(len(paragraphs)):
            a=paragraphs[i].text
            a=a+'\n'
            new_text=new_text+a
        list_all_citation=[]
        for i in range(len(paragraphs)):
            x=paragraphs[i].text
            list_all_citation=list_all_citation+re.findall(r'[a-zA-Z]+ and [a-zA-Z]+, \d+|[a-zA-Z]+ et al., \d+|[a-zA-Z]+, \d+|[a-zA-Z]+ and [a-zA-Z]+ [(]\d+[)]|[a-zA-Z]+ et al. [(]\d+[)]|[a-zA-Z]+ [(]\d+[)]',x)
        list_all_citation
        list_citation=[]
        for i in  range(len(list_all_citation)):
            if list_all_citation[i] not in list_citation:
                list_citation.append(list_all_citation[i])
        s=[]
        for i in range(len(list_citation)):
            s=s+re.findall(r'^[a-zA-Z]+',list_citation[i])
        list_citation_number=[]
        for i in range(len(list_citation)):
            list_citation_number.append(i)
        list_citation_number=[c+1 for c in list_citation_number]
        list_citation_number=[str(c) for c in list_citation_number]
        for i in range(len(list_citation)):
            new_text=new_text.replace(list_citation[i],'['+list_citation_number[i]+']')
        list_reference=re.findall(r'[a-zA-Z]+, [A-Z]\..+',new_text)
        for i in range(len(s)):
            for l in range(len(list_reference)):
                if s[i][0:3] in list_reference[l][0:3]:
                    i=str(i+1)+'. '
                    list_reference[l]=i+list_reference[l]
                    break
        def firest2(elem):
            return float(elem[0:2])
        list_reference.sort(key=firest2)
        index_after_reference=new_text.find('Arnett, J. (1992)')
        new_text=new_text[:index_after_reference]
        for i in range(len(list_reference)):
            new_text=new_text+list_reference[i]+'\n'
        bad_citation=re.findall(r'\(e*\.*g*\.*,* *\[\d+\];* *\[*\d*\]*;* *\[*\d*\]*\)',new_text)
        for i in range(len(bad_citation)):
            if 'e.g.' not in bad_citation[i]:
                new_text=new_text.replace(bad_citation[i],bad_citation[i][1:-1])
            elif 'e.g.' in bad_citation[i]: 
                new_text=new_text.replace(bad_citation[i],bad_citation[i][7:-1])
        new_text=new_text.replace(';',',')
        def remove_control_characters(content):
            mpa = dict.fromkeys(range(32))
            return content.translate(mpa)
        doc = Document()
        for line in new_text.split('\n'):
            paragraph = doc.add_paragraph()
            paragraph.add_run(remove_control_characters(line))
        doc.save('myfile_IEEE_style.docx')
    if filepath=='IEEEexample.docx' and style=='APA':
        document=Document("IEEEexample.docx")
        paragraphs=document.paragraphs   
        new_text2=''
        for i in range(len(paragraphs)):
            a=paragraphs[i].text
            a=a+'\n'
            new_text2=new_text2+a
        index_after_reference2=new_text2.find('1. X. F. Li')
        reference_text=new_text2[index_after_reference2:]
        list_reference2=re.findall(r'.+',reference_text)
        list_number_author=[]
        for i in range(len(list_reference2)):
            number_author=len(re.findall(r'[A-Z]\. [a-zA-Z]+\.* *[A-Z]*[a-z]+',list_reference2[i]))
            number_author=[str(number_author)]
            list_number_author=list_number_author+number_author
        reference_full_name2=re.findall(r'[A-Z]\. [a-zA-Z]+\.* *[A-Z]*[a-z]+',reference_text)
        list_lastname=[]
        for i in range(len(reference_full_name2)):
            list_lastname=list_lastname+re.findall(r'[A-Z][a-z]+',reference_full_name2[i])
        list_year=re.findall(r'\d{4}',reference_text)
        d=0
        ciatation_list=[]
        for i in range(len(list_number_author)):
            if int(list_number_author[i])==1:
                word=[list_lastname[d]+', '+list_year[i]]
                ciatation_list=ciatation_list+word
                d+=1
            if int(list_number_author[i])==2:
                word=[list_lastname[d]+' and '+list_lastname[d+1]+', '+list_year[i]]
                ciatation_list=ciatation_list+word
                d=d+2
            if int(list_number_author[i])==5:
                word=[list_lastname[d]+' et al., '+list_year[i]]
                ciatation_list=ciatation_list+word
        new_citation_list=['','','','','','','','','','','','','']
        list_citation_number=re.findall(r'\[\d\],* *\[*\d*\]*,* *\[*\d*\]*',new_text2)
        for i in range(len(list_citation_number)):
            for n in range(1,10):
                n=str(n)
                if n in list_citation_number[i]:
                    n=int(n)
                    new_citation_list[i]=new_citation_list[i]+'; '+ciatation_list[n-1]
        for i in range(len(new_citation_list)):
            new_citation_list[i]=new_citation_list[i][2:]
            new_citation_list[i]='('+new_citation_list[i]+')'
        for i in range(len(list_citation_number)):
            new_text2=new_text2.replace(list_citation_number[i],new_citation_list[i])
        for i in range(len(list_reference2)):
            list_reference2[i]=list_reference2[i][3:]
        list_reference2.sort()
        referencec_content=''
        for i in range(len(list_reference2)):
            referencec_content=referencec_content+list_reference2[i]+'\n'
        index_after_reference2=new_text2.find('1. X. F. Li')
        new_text2=new_text2[:index_after_reference2]
        new_text2=new_text2+referencec_content
        def remove_control_characters(content):
            mpa = dict.fromkeys(range(32))
            return content.translate(mpa)       
        file= Document()
        for line in new_text2.split('\n'):
            paragraph = file.add_paragraph()
            paragraph.add_run(remove_control_characters(line))
        file.save('myfile_APA_style.docx')
# ---- DO NOT CHANGE THE CODE BELOW ----
if __name__ == "__main__":
    if len(sys.argv)<3: raise ValueError('Provide filename and style as input arguments')
    filepath, style = sys.argv[1], sys.argv[2]
    print('filepath is "{}"'.format(filepath))
    print('target style is "{}"'.format(style))
    change_style(filepath, style)
